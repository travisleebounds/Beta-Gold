from __future__ import annotations
from pathlib import Path
from typing import Any, Dict, List
import csv

def extract_docx(p: Path) -> List[Dict[str, Any]]:
    from docx import Document
    doc = Document(str(p))
    chunks: List[Dict[str, Any]] = []

    texts = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            texts.append(t)
    if texts:
        chunks.append({
            "kind": "text",
            "text": "\n".join(texts),
            "provenance": {"file": p.name, "type": "docx", "locator": "paragraphs"}
        })

    for ti, table in enumerate(doc.tables):
        rows = []
        for r in table.rows:
            rows.append([c.text.strip() for c in r.cells])
        if not rows:
            continue
        headers = rows[0]
        data = rows[1:]
        chunks.append({
            "kind": "table",
            "table": {"headers": headers, "rows": data},
            "provenance": {"file": p.name, "type": "docx", "locator": f"table[{ti}]"}
        })

    return chunks

def extract_xlsx(p: Path) -> List[Dict[str, Any]]:
    import pandas as pd
    chunks: List[Dict[str, Any]] = []
    xls = pd.ExcelFile(p)
    for sheet in xls.sheet_names:
        try:
            df = pd.read_excel(p, sheet_name=sheet, dtype=str)
        except Exception:
            continue
        if df is None or df.empty:
            continue
        df = df.fillna("")
        headers = [str(c) for c in df.columns.tolist()]
        rows = df.head(200).values.tolist()
        chunks.append({
            "kind": "table",
            "table": {"headers": headers, "rows": rows},
            "provenance": {"file": p.name, "type": "xlsx", "locator": f"sheet:{sheet}"}
        })
    return chunks

def extract_csv(p: Path) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    with p.open("r", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        return chunks
    headers = rows[0]
    data = rows[1:201]
    chunks.append({
        "kind": "table",
        "table": {"headers": headers, "rows": data},
        "provenance": {"file": p.name, "type": "csv", "locator": "rows:1-200"}
    })
    return chunks

def extract_pdf_text(p: Path) -> List[Dict[str, Any]]:
    from pypdf import PdfReader
    chunks: List[Dict[str, Any]] = []
    reader = PdfReader(str(p))
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        txt = txt.strip()
        if not txt:
            continue
        chunks.append({
            "kind": "text",
            "text": txt,
            "provenance": {"file": p.name, "type": "pdf", "locator": f"page:{i+1}"}
        })
    return chunks
