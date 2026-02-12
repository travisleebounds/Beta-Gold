from __future__ import annotations
from pathlib import Path
from typing import Any, Dict, List
from datetime import datetime

def write_updated_memo(out_path: Path, latest_facts: List[Dict[str, Any]]):
    from docx import Document

    doc = Document()
    doc.add_heading("UPDATED MEMO — Ingest Summary", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")

    doc.add_heading("Latest Facts", level=2)
    if not latest_facts:
        doc.add_paragraph("No facts extracted.")
    else:
        for f in latest_facts[:200]:
            dt = (f.get("date") or "")[:10]
            metric = f.get("metric", "")
            val = f.get("value", "")
            unit = f.get("unit", "")
            juris = f.get("jurisdiction", "")
            state = f.get("state") or ""
            city = f.get("city") or ""
            desc = f.get("description", "")
            src = f.get("source", {})
            src_txt = f"{src.get('file','')} — {src.get('locator','')}"

            doc.add_paragraph(f"[{dt}] {metric}: {val} {unit} ({juris} {state} {city})")
            if desc:
                doc.add_paragraph(f"  • {desc}")
            doc.add_paragraph(f"  • Source: {src_txt}")

    doc.add_heading("Notes", level=2)
    doc.add_paragraph("Auto-generated. Validate critical items against sources above.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
