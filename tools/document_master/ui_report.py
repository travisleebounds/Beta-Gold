"""
Document Master — Streamlit Report Generator UI
=================================================
Renders the retro terminal report generator inside the dashboard.
Connects to the DocumentMaster engine for actual report generation.

Usage in app.py:
    from tools.document_master.ui_report import render_report_generator
    render_report_generator(member_data, dashboard_context)
"""

import streamlit as st
import json
import time
import threading
import queue
from datetime import datetime


def render_report_generator(member_data: dict, dashboard_context: str = ""):
    """
    Render the report generation UI for a member profile.
    
    Args:
        member_data: dict with keys: id, name, party, area
        dashboard_context: string context from build_dashboard_context()
    """
    member_id = member_data.get("id", "Unknown")
    member_name = member_data.get("name", "Unknown Member")
    
    st.markdown("---")
    st.markdown("### 🖥️ Document Master — Report Generator")
    
    # Check engine status
    try:
        from tools.document_master.engine import DocumentMaster
        dm = DocumentMaster()
        status = dm.status()
        
        if status["total_chunks"] > 0:
            st.caption(f"📚 {status['documents_indexed']} docs indexed | {status['total_chunks']} chunks | Model: {status['model']}")
        else:
            st.caption(f"⚠️ No documents indexed yet. Go to 📥 Ingested Docs to upload files.")
        
        if not status["ollama_running"]:
            st.warning("⚠️ Ollama is not running. Start it with: `systemctl start ollama`")
            return
            
    except ImportError:
        st.error("❌ Document Master engine not installed. Run: `bash setup_document_master.sh`")
        return
    except Exception as e:
        st.warning(f"⚠️ Document Master status check failed: {e}")
        st.caption("Reports may still work if Ollama is running.")
    
    # Report type selection
    col1, col2 = st.columns(2)
    
    with col1:
        brief_clicked = st.button(
            "📋 Policy Brief",
            help="1-2 page executive summary with key findings",
            key=f"brief_{member_id}",
            use_container_width=True,
        )
    
    with col2:
        nuke_clicked = st.button(
            "☢️ Data Nuke",
            help="10+ page comprehensive deep dive — everything we know",
            key=f"nuke_{member_id}",
            use_container_width=True,
            type="primary",
        )
    
    report_type = None
    if brief_clicked:
        report_type = "brief"
    elif nuke_clicked:
        report_type = "nuke"
    
    if report_type:
        _run_report_generation(dm, member_data, report_type, dashboard_context)


def _run_report_generation(dm, member_data: dict, report_type: str, dashboard_context: str):
    """Run the report generation with streaming terminal output."""
    
    member_id = member_data.get("id", "Unknown")
    member_name = member_data.get("name", "Unknown Member")
    is_nuke = report_type == "nuke"
    
    st.markdown("---")
    
    # Terminal-style header
    if is_nuke:
        st.markdown(
            f"""<div style="background: #1a0000; border: 1px solid #3a0000; border-radius: 6px; 
            padding: 12px 16px; font-family: monospace; color: #ff2020; font-size: 13px;">
            ☢️ DATA NUKE — {member_id} — {member_name}<br>
            <span style="color: #666;">Generating comprehensive report... this may take 2-3 minutes.</span>
            </div>""",
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f"""<div style="background: #001a00; border: 1px solid #003a00; border-radius: 6px; 
            padding: 12px 16px; font-family: monospace; color: #00ff41; font-size: 13px;">
            📋 POLICY BRIEF — {member_id} — {member_name}<br>
            <span style="color: #666;">Generating executive summary...</span>
            </div>""",
            unsafe_allow_html=True,
        )
    
    # Progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Terminal output area
    terminal = st.empty()
    
    # Generate
    full_report = ""
    log_lines = []
    timestamp = datetime.now().strftime("%H:%M:%S")
    
    try:
        for event in dm.generate_report_stream(member_data, report_type, dashboard_context):
            
            if "stage" in event:
                progress = event["progress"] / 100
                progress_bar.progress(min(progress, 0.99))
                status_text.caption(f"⏳ {event['stage']}...")
                log_lines.append(f"[{timestamp}] > {event['stage']}... ✓")
                
                # Show logs so far
                terminal.code("\n".join(log_lines), language="bash")
                time.sleep(0.1)
            
            elif "token" in event:
                full_report += event["token"]
                # Update terminal with streaming report
                accent = "☢️" if is_nuke else "📋"
                terminal.code(
                    "\n".join(log_lines) + f"\n\n{accent} REPORT OUTPUT:\n" + full_report,
                    language="text",
                )
            
            elif "done" in event:
                progress_bar.progress(1.0)
                status_text.caption(f"✅ Report complete — {event.get('sources', 0)} source documents referenced")
            
            elif "error" in event:
                st.error(f"❌ {event['error']}")
                return
    
    except Exception as e:
        st.error(f"❌ Report generation failed: {e}")
        return
    
    if full_report:
        st.markdown("---")
        
        # Download buttons
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                "📄 Download as TXT",
                data=full_report,
                file_name=f"report_{member_id}_{report_type}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                key=f"dl_txt_{member_id}_{report_type}",
            )
        
        with col2:
            # Generate DOCX if python-docx is available
            try:
                from docx import Document as DocxDocument
                from io import BytesIO
                
                doc = DocxDocument()
                doc.add_heading(f"{'Policy Brief' if report_type == 'brief' else 'Data Nuke Report'} — {member_id}", 0)
                doc.add_paragraph(f"Member: {member_data.get('name', 'Unknown')}")
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_paragraph("")
                
                for line in full_report.split("\n"):
                    doc.add_paragraph(line)
                
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                
                st.download_button(
                    "📝 Download as DOCX",
                    data=buf.getvalue(),
                    file_name=f"report_{member_id}_{report_type}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_docx_{member_id}_{report_type}",
                )
            except ImportError:
                st.caption("Install python-docx for DOCX export")
        
        with col3:
            st.download_button(
                "📊 Download as JSON",
                data=json.dumps({
                    "report_type": report_type,
                    "member": member_data,
                    "generated_at": datetime.now().isoformat(),
                    "content": full_report,
                }, indent=2),
                file_name=f"report_{member_id}_{report_type}_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                key=f"dl_json_{member_id}_{report_type}",
            )
