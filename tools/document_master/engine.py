"""
Document Master Engine — IDOT Dashboard
========================================
Local AI document engine powered by Ollama + ChromaDB.

Handles:
  - Tiered document ingestion:
    • GOLD tier: Current standard reports/policies (highest search priority)
    • ARCHIVE tier: Historical corpus (searchable background knowledge)
  - Vector embedding & storage in ChromaDB
  - Priority-weighted semantic search
  - Batch ingestion with resume (handles 60K+ docs)
  - Report generation (Policy Brief / Data Nuke) via Ollama streaming

Usage:
  from tools.document_master.engine import DocumentMaster
  dm = DocumentMaster()
  
  # Ingest gold standard docs (priority)
  dm.ingest_directory("/path/to/gold_reports", tier="gold")
  
  # Ingest historical archive (batch, resumable)
  dm.batch_ingest("/path/to/archive", tier="archive")
  
  # Search (gold docs weighted higher)
  results = dm.search("federal funding allocations")
  
  # Generate reports
  report = dm.generate_report(member_data, report_type="brief")
"""

import os
import json
import hashlib
import time
import logging
import traceback
from pathlib import Path
from datetime import datetime
from typing import Generator

# Document parsing
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

import csv

# Vector store
import chromadb
from chromadb.config import Settings as ChromaSettings

# Text splitting
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Ollama client
import ollama as ollama_client

logger = logging.getLogger("document_master")
logging.basicConfig(level=logging.INFO)

# ═══════════════════════════════════════════════════════════════
# Configuration
# ═══════════════════════════════════════════════════════════════

DEFAULT_MODEL = os.environ.get("DOCMASTER_MODEL", "qwen2.5-coder:7b")
FALLBACK_MODEL = "llama3.1:8b"
CHROMA_DIR = os.environ.get("DOCMASTER_CHROMA_DIR", "data/vectorstore")
COLLECTION_NAME = "idot_documents"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
TOP_K_RESULTS = 15

TIER_WEIGHTS = {
    "gold": 2.0,
    "standard": 1.0,
    "archive": 0.7,
}

BATCH_PROGRESS_FILE = "data/ingest/batch_progress.json"
BATCH_LOG_FILE = "logs/ingest.log"
BATCH_SAVE_INTERVAL = 50


# ═══════════════════════════════════════════════════════════════
# Document Parsers
# ═══════════════════════════════════════════════════════════════

def parse_pdf(filepath: str) -> str:
    if PdfReader is None:
        raise ImportError("PyPDF2 not installed")
    reader = PdfReader(filepath)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def parse_docx(filepath: str) -> str:
    if DocxDocument is None:
        raise ImportError("python-docx not installed")
    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def parse_xlsx(filepath: str) -> str:
    if openpyxl is None:
        raise ImportError("openpyxl not installed")
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    content = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        content.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                content.append(" | ".join(cells))
    wb.close()
    return "\n".join(content)


def parse_csv_file(filepath: str) -> str:
    content = []
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            if any(cell.strip() for cell in row):
                content.append(" | ".join(row))
    return "\n".join(content)


def parse_text(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".csv": parse_csv_file,
    ".tsv": parse_csv_file,
    ".txt": parse_text,
    ".md": parse_text,
    ".json": parse_text,
}


def parse_file(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(filepath)


def file_hash(filepath: str) -> str:
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _collect_files(directory: str, recursive: bool = True) -> list:
    directory = Path(directory)
    files = []
    if recursive:
        for root, dirs, filenames in os.walk(str(directory)):
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            for fname in sorted(filenames):
                fpath = Path(root) / fname
                if fpath.suffix.lower() in PARSERS:
                    files.append(str(fpath))
    else:
        for fpath in sorted(directory.iterdir()):
            if fpath.is_file() and fpath.suffix.lower() in PARSERS:
                files.append(str(fpath))
    return files


# ═══════════════════════════════════════════════════════════════
# Document Master Engine
# ═══════════════════════════════════════════════════════════════

class DocumentMaster:

    def __init__(self, chroma_dir: str = CHROMA_DIR, model: str = DEFAULT_MODEL):
        self.model = model
        self.chroma_dir = chroma_dir
        
        os.makedirs(chroma_dir, exist_ok=True)
        self.chroma_client = chromadb.PersistentClient(path=chroma_dir)
        self.collection = self.chroma_client.get_or_create_collection(
            name=COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"}
        )
        
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        
        self.index_path = Path("data/ingest/docmaster_index.json")
        self.index = self._load_index()
        
        logger.info(f"DocumentMaster initialized: model={model}, docs={self.collection.count()}")

    def _load_index(self) -> dict:
        if self.index_path.exists():
            with open(self.index_path) as f:
                return json.load(f)
        return {"documents": {}, "last_updated": None, "stats": {"gold": 0, "archive": 0, "standard": 0}}

    def _save_index(self):
        self.index["last_updated"] = datetime.now().isoformat()
        self.index_path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.index_path, "w") as f:
            json.dump(self.index, f, indent=2)

    def _check_ollama(self) -> bool:
        try:
            models = ollama_client.list()
            available = [m.get("name", m.get("model", "")) for m in models.get("models", [])]
            for m in available:
                if self.model.split(":")[0] in m:
                    return True
            return False
        except Exception:
            return False

    # ─── Ingestion ──────────────────────────────────────────────

    def ingest_file(self, filepath: str, tier: str = "standard", force: bool = False) -> dict:
        filepath = str(Path(filepath).resolve())
        fname = Path(filepath).name
        
        try:
            fhash = file_hash(filepath)
        except Exception as e:
            return {"file": fname, "status": "error", "reason": f"Cannot read: {e}"}
        
        if not force and fname in self.index["documents"]:
            existing = self.index["documents"][fname]
            if existing.get("sha256") == fhash and existing.get("tier") == tier:
                return {"file": fname, "status": "skipped", "reason": "already ingested"}
        
        try:
            text = parse_file(filepath)
        except Exception as e:
            return {"file": fname, "status": "error", "reason": str(e)}
        
        if not text.strip():
            return {"file": fname, "status": "error", "reason": "empty content"}
        
        chunks = self.splitter.split_text(text)
        if not chunks:
            return {"file": fname, "status": "error", "reason": "no chunks"}
        
        # Remove old entries
        try:
            existing = self.collection.get(where={"source_file": fname})
            if existing and existing["ids"]:
                self.collection.delete(ids=existing["ids"])
        except Exception:
            pass
        
        ids = [f"{fname}__chunk_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "source_file": fname,
                "source_path": filepath,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "sha256": fhash,
                "tier": tier,
                "ingested_at": datetime.now().isoformat(),
            }
            for i in range(len(chunks))
        ]
        
        batch_size = 100
        for start in range(0, len(chunks), batch_size):
            end = min(start + batch_size, len(chunks))
            self.collection.add(
                ids=ids[start:end],
                documents=chunks[start:end],
                metadatas=metadatas[start:end],
            )
        
        self.index["documents"][fname] = {
            "sha256": fhash,
            "path": filepath,
            "chunks": len(chunks),
            "chars": len(text),
            "tier": tier,
            "ingested_at": datetime.now().isoformat(),
        }
        
        if "stats" not in self.index:
            self.index["stats"] = {"gold": 0, "archive": 0, "standard": 0}
        self.index["stats"][tier] = self.index["stats"].get(tier, 0) + 1
        self._save_index()
        
        return {"file": fname, "status": "ingested", "chunks": len(chunks), "chars": len(text), "tier": tier}

    def ingest_directory(self, directory: str, tier: str = "standard",
                          force: bool = False, recursive: bool = False) -> list:
        results = []
        files = _collect_files(directory, recursive=recursive)
        if not files:
            return [{"error": f"No supported files in {directory}"}]
        for filepath in files:
            result = self.ingest_file(filepath, tier=tier, force=force)
            results.append(result)
        return results

    # ─── Batch Ingestion ────────────────────────────────────────

    def batch_ingest(self, directory: str, tier: str = "archive",
                      recursive: bool = True, force: bool = False,
                      callback=None) -> dict:
        logger.info(f"Scanning {directory}...")
        all_files = _collect_files(directory, recursive=recursive)
        total = len(all_files)
        logger.info(f"Found {total} supported files")
        
        if total == 0:
            return {"status": "empty", "total": 0}
        
        progress_path = Path(BATCH_PROGRESS_FILE)
        progress_path.parent.mkdir(parents=True, exist_ok=True)
        
        completed = set()
        if progress_path.exists() and not force:
            try:
                completed = set(json.load(open(progress_path)).get("completed", []))
                logger.info(f"Resuming: {len(completed)}/{total} already done")
            except Exception:
                pass
        
        log_path = Path(BATCH_LOG_FILE)
        log_path.parent.mkdir(parents=True, exist_ok=True)
        fh = logging.FileHandler(str(log_path))
        fh.setFormatter(logging.Formatter("%(asctime)s %(message)s"))
        logger.addHandler(fh)
        
        stats = {"total": total, "ingested": 0, "skipped": 0, "errors": 0,
                 "resumed_from": len(completed), "start_time": datetime.now().isoformat()}
        error_files = []
        
        for i, filepath in enumerate(all_files):
            fname = Path(filepath).name
            
            if fname in completed and not force:
                stats["skipped"] += 1
                continue
            
            try:
                result = self.ingest_file(filepath, tier=tier, force=force)
                if result.get("status") == "ingested":
                    stats["ingested"] += 1
                elif result.get("status") == "skipped":
                    stats["skipped"] += 1
                else:
                    stats["errors"] += 1
                    error_files.append({"file": fname, "reason": result.get("reason", "?")})
                completed.add(fname)
            except Exception as e:
                stats["errors"] += 1
                error_files.append({"file": fname, "reason": str(e)})
                completed.add(fname)
            
            processed = stats["ingested"] + stats["skipped"] + stats["errors"]
            if callback:
                callback(processed, total, fname)
            
            if processed % BATCH_SAVE_INTERVAL == 0:
                with open(progress_path, "w") as f:
                    json.dump({"completed": list(completed), "stats": stats,
                               "last_saved": datetime.now().isoformat()}, f)
                self._save_index()
                
                elapsed = max(1, (datetime.now() - datetime.fromisoformat(stats["start_time"])).seconds)
                rate = processed / elapsed
                remaining = (total - processed) / max(rate, 0.01)
                logger.info(f"Progress: {processed}/{total} ({stats['ingested']} new, {stats['errors']} err) ~{remaining/60:.0f}m left")
        
        stats["end_time"] = datetime.now().isoformat()
        stats["error_files"] = error_files[:100]
        
        with open(progress_path, "w") as f:
            json.dump({"completed": list(completed), "stats": stats, "finished": True}, f)
        self._save_index()
        
        logger.removeHandler(fh)
        fh.close()
        
        return stats

    # ─── Search ─────────────────────────────────────────────────

    def search(self, query: str, n_results: int = TOP_K_RESULTS,
               filter_file: str = None, tier: str = None,
               gold_boost: bool = True) -> list:
        if self.collection.count() == 0:
            return []
        
        where = None
        if filter_file and tier:
            where = {"$and": [{"source_file": filter_file}, {"tier": tier}]}
        elif filter_file:
            where = {"source_file": filter_file}
        elif tier:
            where = {"tier": tier}
        
        fetch_n = min(n_results * 3 if gold_boost else n_results, self.collection.count())
        
        results = self.collection.query(
            query_texts=[query],
            n_results=fetch_n,
            where=where,
        )
        
        hits = []
        if results and results["documents"]:
            for i, doc in enumerate(results["documents"][0]):
                meta = results["metadatas"][0][i] if results["metadatas"] else {}
                distance = results["distances"][0][i] if results["distances"] else 0
                similarity = 1 - distance
                doc_tier = meta.get("tier", "standard")
                weight = TIER_WEIGHTS.get(doc_tier, 1.0) if gold_boost else 1.0
                
                hits.append({
                    "text": doc,
                    "source_file": meta.get("source_file", "unknown"),
                    "source_path": meta.get("source_path", ""),
                    "chunk_index": meta.get("chunk_index", 0),
                    "tier": doc_tier,
                    "score": similarity * weight,
                    "raw_score": similarity,
                })
        
        hits.sort(key=lambda x: x["score"], reverse=True)
        return hits[:n_results]

    # ─── Report Generation ──────────────────────────────────────

    def _build_report_prompt(self, member_data: dict, report_type: str,
                              context_chunks: list, dashboard_context: str = "") -> str:
        member_id = member_data.get("id", "Unknown")
        member_name = member_data.get("name", "Unknown Member")
        party = member_data.get("party", "?")
        area = member_data.get("area", "Illinois")
        
        doc_lines = []
        for c in context_chunks:
            tier_marker = "⭐ GOLD" if c.get("tier") == "gold" else "📁"
            doc_lines.append(f"[{tier_marker} | {c['source_file']}]\n{c['text']}")
        doc_context = "\n\n---\n\n".join(doc_lines)
        
        if report_type == "brief":
            return f"""You are Document Master, an AI report generator for the Illinois Department of Transportation Dashboard.

Generate a POLICY BRIEF (1-2 pages) for the following member.

MEMBER INFORMATION:
- ID: {member_id}
- Name: {member_name}
- Party: {party}
- Area: {area}

DASHBOARD DATA:
{dashboard_context[:3000]}

RELEVANT DOCUMENTS (⭐ GOLD = current office standard, 📁 = archive):
{doc_context[:4000]}

INSTRUCTIONS:
Generate a concise policy brief with these sections:
1. EXECUTIVE SUMMARY — 2-3 sentence overview
2. KEY FINDINGS — bullet points of most important facts
3. POLICY REFERENCES — relevant policies, bills, and compliance status
4. RECOMMENDATION — 1-2 sentence action item

IMPORTANT: When gold-standard documents are available, follow their format and style closely. They represent the current office standard.

Format as a clean text report with clear section headers.
Use ═ and ─ characters for borders. Include the member name and date at the top.
Be specific — cite actual data from the context. If data is missing, note "Data pending".
"""
        else:
            return f"""You are Document Master, an AI report generator for the Illinois Department of Transportation Dashboard.

Generate a COMPREHENSIVE DATA NUKE REPORT (10+ pages) for the following member.

MEMBER INFORMATION:
- ID: {member_id}
- Name: {member_name}
- Party: {party}
- Area: {area}

DASHBOARD DATA:
{dashboard_context[:5000]}

RELEVANT DOCUMENTS (⭐ GOLD = current office standard, 📁 = archive):
{doc_context[:8000]}

INSTRUCTIONS:
Generate an exhaustive report with ALL sections:
1. EXECUTIVE SUMMARY — Full overview with confidence scores
2. MEMBER PROFILE & HISTORY — Complete background
3. POLICY COMPLIANCE AUDIT — Every relevant policy checked
4. FEDERAL FUNDING ANALYSIS — Formula allocations, grants, per-capita comparisons
5. TRANSPORTATION INFRASTRUCTURE — Road events, construction, closures
6. LEGISLATIVE ACTIVITY — Bills sponsored, committee work
7. RISK ASSESSMENT MATRIX — Rate each area: LOW / MEDIUM / HIGH
8. COMPARATIVE ANALYSIS — How this member compares to peers
9. HISTORICAL TIMELINE — Key events chronologically
10. DOCUMENT CROSS-REFERENCE — Which source docs informed each section
11. RECOMMENDATIONS & ACTION ITEMS — Specific next steps
12. APPENDIX — Source document list with dates

IMPORTANT: Follow gold-standard document format and style when available.

Format with ══ double borders for major sections, ── for subsections.
[■] completed items, [□] pending. Include TABLE OF CONTENTS.
Be exhaustive. If data is missing, note "DATA PENDING — requires [source]".
"""

    def generate_report_stream(self, member_data: dict, report_type: str = "brief",
                                dashboard_context: str = "") -> Generator[dict, None, None]:
        stages = [
            "Connecting to Document Master",
            "Loading member profile data",
            "Querying gold-standard policy database",
            "Searching document archive",
        ]
        if report_type == "nuke":
            stages.extend([
                "Cross-referencing compliance records",
                "Analyzing historical data",
                "Processing all related documents",
                "Building risk assessment matrix",
            ])
        stages.append("Generating report")
        
        member_id = member_data.get("id", "")
        member_name = member_data.get("name", "Unknown")
        
        for i, stage in enumerate(stages[:-1]):
            yield {"stage": stage, "progress": (i / len(stages)) * 100}
            time.sleep(0.3)
        
        search_queries = [
            f"{member_name} {member_data.get('area', '')}",
            f"{member_id} transportation funding",
            f"{member_id} policy compliance",
            "Illinois transportation federal funding",
        ]
        if report_type == "nuke":
            search_queries.extend([
                f"{member_id} grants discretionary",
                f"{member_id} road construction closures",
                f"{member_data.get('area', '')} infrastructure",
                "IIJA formula allocations Illinois",
            ])
        
        all_chunks = []
        seen_texts = set()
        for query in search_queries:
            results = self.search(query, n_results=5, gold_boost=True)
            for r in results:
                if r["text"] not in seen_texts:
                    all_chunks.append(r)
                    seen_texts.add(r["text"])
        
        prompt = self._build_report_prompt(member_data, report_type, all_chunks, dashboard_context)
        
        yield {"stage": "Generating report", "progress": 90}
        
        if not self._check_ollama():
            yield {"error": f"Ollama not available. Run: ollama pull {self.model}"}
            return
        
        full_text = ""
        try:
            stream = ollama_client.chat(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                stream=True,
                options={"temperature": 0.3, "num_predict": 4096 if report_type == "brief" else 8192, "top_p": 0.9},
            )
            for chunk in stream:
                token = chunk.get("message", {}).get("content", "")
                if token:
                    full_text += token
                    yield {"token": token}
            
            gold_count = sum(1 for c in all_chunks if c.get("tier") == "gold")
            yield {"done": True, "full_text": full_text, "sources": len(all_chunks), "gold_sources": gold_count}
        except Exception as e:
            yield {"error": f"Generation failed: {e}"}

    def generate_report(self, member_data: dict, report_type: str = "brief",
                         dashboard_context: str = "") -> str:
        full_text = ""
        for event in self.generate_report_stream(member_data, report_type, dashboard_context):
            if "token" in event:
                full_text += event["token"]
            elif "error" in event:
                raise RuntimeError(event["error"])
        return full_text

    # ─── Status ─────────────────────────────────────────────────

    def status(self) -> dict:
        ollama_ok = self._check_ollama()
        tier_counts = {"gold": 0, "standard": 0, "archive": 0}
        for doc_info in self.index.get("documents", {}).values():
            t = doc_info.get("tier", "standard")
            tier_counts[t] = tier_counts.get(t, 0) + 1
        
        return {
            "ollama_running": ollama_ok,
            "model": self.model,
            "documents_indexed": len(self.index.get("documents", {})),
            "total_chunks": self.collection.count(),
            "tier_counts": tier_counts,
            "chroma_dir": self.chroma_dir,
            "last_updated": self.index.get("last_updated"),
        }


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys
    
    dm = DocumentMaster()
    
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python -m tools.document_master.engine status")
        print("  python -m tools.document_master.engine ingest <path> [--tier gold|archive|standard]")
        print("  python -m tools.document_master.engine batch <path> [--tier gold|archive]")
        print("  python -m tools.document_master.engine search <query>")
        print("  python -m tools.document_master.engine report <member_id> [brief|nuke]")
        sys.exit(1)
    
    cmd = sys.argv[1]
    
    if cmd == "status":
        print(json.dumps(dm.status(), indent=2))
    
    elif cmd == "ingest":
        path = sys.argv[2] if len(sys.argv) > 2 else "ingest_inbox"
        tier = "standard"
        if "--tier" in sys.argv:
            tier = sys.argv[sys.argv.index("--tier") + 1]
        p = Path(path)
        if p.is_dir():
            results = dm.ingest_directory(str(p), tier=tier)
        else:
            results = [dm.ingest_file(str(p), tier=tier)]
        for r in results:
            s = r.get("status", "?")
            icon = "✅" if s == "ingested" else "⏭️" if s == "skipped" else "❌"
            print(f"  {icon} {r.get('file','?')}: {s} ({r.get('chunks', 0)} chunks) [{r.get('tier', '?')}]")
        print(f"\nTotal chunks: {dm.collection.count()}")
    
    elif cmd == "batch":
        path = sys.argv[2] if len(sys.argv) > 2 else "."
        tier = "archive"
        if "--tier" in sys.argv:
            tier = sys.argv[sys.argv.index("--tier") + 1]
        
        def progress_cb(processed, total, fname):
            if processed % 100 == 0 or processed == total:
                print(f"  [{(processed/total)*100:5.1f}%] {processed}/{total} — {fname}")
        
        print(f"Batch ingesting {path} as tier={tier}...")
        stats = dm.batch_ingest(path, tier=tier, callback=progress_cb)
        print(f"\nDone: {stats.get('ingested',0)} ingested, {stats.get('skipped',0)} skipped, {stats.get('errors',0)} errors")
        print(f"Total chunks: {dm.collection.count()}")
    
    elif cmd == "search":
        query = " ".join(sys.argv[2:])
        for r in dm.search(query):
            icon = "⭐" if r["tier"] == "gold" else "📁"
            print(f"  {icon} [{r['source_file']}] (score: {r['score']:.3f}, tier: {r['tier']})")
            print(f"    {r['text'][:120]}...\n")
    
    elif cmd == "report":
        mid = sys.argv[2] if len(sys.argv) > 2 else "IL-01"
        rtype = sys.argv[3] if len(sys.argv) > 3 else "brief"
        member_data = {"id": mid, "name": "Test Member", "party": "D", "area": "Illinois"}
        for event in dm.generate_report_stream(member_data, rtype):
            if "stage" in event:
                print(f"  [{event['progress']:.0f}%] {event['stage']}...")
            elif "token" in event:
                print(event["token"], end="", flush=True)
            elif "done" in event:
                print(f"\n\n--- Complete ({event['sources']} sources, {event.get('gold_sources',0)} gold) ---")
            elif "error" in event:
                print(f"\n  ❌ {event['error']}")
